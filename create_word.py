import pandas as pd
from docx import Document
import random
from docx2pdf import convert
from decimal import Decimal
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os
from docx.oxml.ns import qn
import logging

# 配置日志
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')


def fill_template(template_path, data_df, output_doc_path):
    try:
        doc = Document(template_path)
        original_styles = {}

        # 提取原始文档的样式
        for style in doc.styles:
            original_styles[style.name] = style

        # 随机选择一行数据
        random_row = data_df.sample(n=1).iloc[0]
        #logging.info("读取的数据如下：")
        #logging.info(random_row.to_csv(sep='\t', na_rep='nan'))

        # 处理金额（小写）保留2位小数
        amount_lower = '{:,.2f}'.format(Decimal(random_row["金额"]))

        # 定义需要替换的字段及其对应的值
        replacements = {
            "$产品名称$": str(random_row["产品名称"]).strip(),
            "$基金账号$": str(random_row["基金账号"]).strip(),
            "$证券代码$": str(random_row["证券代码"]).strip(),
            "$证券名称$": str(random_row["证券名称"]).strip(),
            "$金额$": amount_lower,
            "$份额$": '{:,.2f}'.format(Decimal(random_row["份额"])),
            "$基金净值$": str(random_row["基金净值"]).strip()
        }

        def process_para(para):
            """
            处理段落，替换占位符并保持原始样式
            :param para: 段落对象
            """
            para_text = ''.join([run.text for run in para.runs if run.text])
            for key, value in replacements.items():
                if key in para_text:
                    logging.info(f"找到段落中的占位符 {key}，将替换为 {value}")
                    start_index = 0
                    while key in para_text[start_index:]:
                        index = para_text[start_index:].find(key) + start_index
                        run_index = 0
                        pos = 0
                        # 找到占位符起始的 run
                        for i, run in enumerate(para.runs):
                            if pos + len(run.text) > index:
                                run_index = i
                                break
                            pos += len(run.text)

                        start_in_run = index - pos
                        remaining = key
                        temp_runs = []
                        while remaining:
                            current_run = para.runs[run_index]
                            end_in_run = min(start_in_run + len(remaining), len(current_run.text))
                            if end_in_run == len(current_run.text):
                                part_text = current_run.text[:start_in_run] + remaining
                                new_part_text = part_text.replace(key, value)
                                temp_run = para._element._new_r()
                                temp_run.text = new_part_text
                                temp_runs.append(temp_run)
                                remaining = ""
                            else:
                                part_text = current_run.text[:start_in_run] + remaining[:end_in_run - start_in_run]
                                temp_run = para._element._new_r()
                                temp_run.text = part_text
                                temp_runs.append(temp_run)
                                remaining = remaining[end_in_run - start_in_run:]
                                run_index += 1
                                start_in_run = 0

                        # 移除旧的 runs 并添加新的 runs
                        for i in range(len(para.runs)):
                            para._element.remove(para.runs[0]._element)
                        for run in temp_runs:
                            para._element.append(run)

                        # 替换占位符
                        para_text = para_text.replace(key, value, 1)
                        start_index = index + len(value)

                    # 重新设置样式
                    for run in para.runs:
                        original_font_name = run.font.name
                        original_font_size = run.font.size
                        if original_font_name:
                            run.font.name = original_font_name
                        if original_font_size:
                            run.font.size = original_font_size
                        if run.style.name in original_styles:
                            run.style = original_styles[run.style.name]

        # 遍历文档中的所有段落
        for para in doc.paragraphs:
            process_para(para)

        # 遍历文档中的所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        process_para(para)

        # 处理文本框内容
        def process_textbox(shape):
            logging.info("开始处理文本框")
            for p in shape.iter(qn('w:p')):
                # 模拟段落对象
                class MockPara:
                    def __init__(self, p):
                        self._element = p
                        self.runs = []
                        for r in p.iter(qn('w:r')):
                            self.runs.append(r)

                    @property
                    def text(self):
                        return ''.join([r.text for r in self.runs if r.text])

                para = MockPara(p)
                process_para(para)

        for element in doc.element.body:
            if element.tag.endswith('}drawing'):  # 包含文本框的绘图元素
                for shape in element.iter():
                    if shape.tag.endswith('}txbx'):  # 文本框
                        process_textbox(shape)

        # 保存填充后的文档
        doc.save(output_doc_path)
        logging.info(f"成功保存填充后的文档: {output_doc_path}")
    except Exception as e:
        logging.error(f"填充模板 {template_path} 时发生错误: {e}")



if __name__ == "__main__":
    # 表格文件路径
    data_file = "values.xlsx"
    # 模板文件路径
    template_folder = "input"
    output_folder = "output"
    # 确保输出文件夹存在
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    try:
        data = pd.read_excel(data_file, dtype=str)
        # 对数据进行随机排序
        data = data.sample(frac=1).reset_index(drop=True)
        for template_name in os.listdir(template_folder):
            if template_name.endswith('.docx'):
                template_file = os.path.join(template_folder, template_name)
                template_name_pre = template_name.split("模板")[0]
                for i in range(1, 4):
                    output_doc_file = os.path.join(output_folder, f"{template_name_pre}0{i}.docx")
                    # output_pdf_file = os.path.join(output_folder, f"{template_name_pre}0{i}.pdf")
                    fill_template(template_file, data, output_doc_file)
                    # convert(output_doc_file, output_pdf_file)
    except FileNotFoundError as e:
        logging.error(f"错误: 表格文件或模板文件未找到，请检查文件路径。 {e}")
    except Exception as e:
        logging.error(f"错误: 发生了一个未知错误: {e}")
